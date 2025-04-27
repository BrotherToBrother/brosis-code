import customtkinter as ctk
from tkinter import filedialog, messagebox
import os
import threading
import requests
from docx import Document
import datetime
import ollama
from deep_translator import GoogleTranslator

# Setup customtkinter appearance
ctk.set_appearance_mode("System")  # "System", "Dark", "Light"
ctk.set_default_color_theme("blue")  # Color theme: "blue", "green", "dark-blue"

# System dictionary
system = {}

# Fetch installed Ollama models dynamically
def get_installed_ollama_models():
    try:
        response = requests.get('http://localhost:11434/api/tags')
        if response.status_code == 200:
            models = response.json().get('models', [])
            return [model['name'] for model in models]
    except Exception as e:
        print(f"Error fetching models: {e}")
    return []

# Check if model is pulled
def is_model_pulled(model_name):
    try:
        response = requests.get('http://localhost:11434/api/tags')
        if response.status_code == 200:
            models = response.json().get('models', [])
            for model in models:
                if model['name'].lower().startswith(model_name.lower()):
                    return True
        return False
    except Exception as e:
        print(f"Error checking pulled models: {e}")
        return False

# Pull model automatically if missing
def pull_model(model_name):
    try:
        print(f"🔄 Pulling model {model_name} ...")
        response = requests.post('http://localhost:11434/api/pull', json={"name": model_name})
        for line in response.iter_lines():
            if line:
                decoded = line.decode('utf-8')
                print(decoded)
        print(f"✅ Model {model_name} pulled successfully!")
    except Exception as e:
        print(f"Error pulling model: {e}")

# Translate text through multiple languages (decoherence)
def multi_hop_translate(text):
    try:
        translated = GoogleTranslator(source='english', target='arabic').translate(text=text)
        translated = GoogleTranslator(source='arabic', target='zh-CN').translate(text=translated)
        translated = GoogleTranslator(source='zh-CN', target='ru').translate(text=translated)
        translated_back = GoogleTranslator(source='ru', target='english').translate(text=translated)
        return translated_back
    except Exception as e:
        print(f"Translation error: {e}")
        return text

# Setup Ollama model
def setup_ollama(model):
    system['model'] = model
    system['humanizer'] = """You are HUMIZER and DECOHERE, an AI designed to humanize and decohere AI-generated content so that it passes as human-written text in an academic voice. Do not answer questions or modify meanings. Just rewrite naturally."""
    system['grammar'] = "Only fix grammar problems without changing the structure or original meaning."

# Humanizer function
def super_humanize(text):
    try:
        response = ollama.generate(
            model=system['model'],
            raw=False,
            stream=False,
            system=system['humanizer'],
            prompt=text,
            options={'temperature': 2.7}
        )
        return response['response']
    except Exception as e:
        print(f"Humanizing error: {e}")
        return text

# Grammar corrector
def grammar_correct(text):
    try:
        response = ollama.generate(
            model=system['model'],
            raw=False,
            stream=False,
            system=system['grammar'],
            prompt=text,
            options={'temperature': 1.0}
        )
        return response['response']
    except Exception as e:
        print(f"Grammar correction error: {e}")
        return text

# Process a single DOCX file
def process_single_docx(file_path, output_folder, grammar, log_area):
    doc = Document(file_path)
    new_doc = Document()

    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "":
            continue

        output_text = paragraph.text
        output_text = multi_hop_translate(output_text)
        output_text = super_humanize(output_text)

        if grammar:
            output_text = grammar_correct(output_text)

        new_doc.add_paragraph(output_text)

    filename = os.path.basename(file_path)
    name, ext = os.path.splitext(filename)
    timestamp = str(int(datetime.datetime.now().timestamp()))
    output_filename = f"{name}_humanized_{timestamp}{ext}"
    output_path = os.path.join(output_folder, output_filename)
    new_doc.save(output_path)

    log_area.insert("end", f"✅ Processed and saved: {output_filename}\n")
    log_area.see("end")

# Batch processing
def batch_process(input_folder, output_folder, model, grammar, log_area, progress_bar):
    setup_ollama(model)

    docx_files = [f for f in os.listdir(input_folder) if f.endswith(".docx")]
    total_files = len(docx_files)
    processed_files = 0

    for filename in docx_files:
        file_path = os.path.join(input_folder, filename)
        process_single_docx(file_path, output_folder, grammar, log_area)

        processed_files += 1
        progress = processed_files / total_files
        progress_bar.set(progress)

    log_area.insert("end", "\n🎉 Batch processing complete!\n")
    log_area.see("end")
    messagebox.showinfo("Done", "Batch processing complete!")

# Start batch function
def start_batch():
    input_folder = input_folder_var.get()
    output_folder = output_folder_var.get()
    model = model_var.get()
    grammar = grammar_var.get()

    if not input_folder or not output_folder:
        messagebox.showerror("Error", "Please select both input and output folders.")
        return

    log_area.delete("1.0", "end")
    progress_bar.set(0)

    if not is_model_pulled(model):
        answer = messagebox.askyesno("Model Missing", f"The model '{model}' is not pulled yet.\nDo you want to pull it automatically?")
        if answer:
            pull_model(model)
        else:
            messagebox.showerror("Error", "Model is not available, cannot continue.")
            return

    threading.Thread(target=batch_process, args=(input_folder, output_folder, model, grammar, log_area, progress_bar)).start()

# GUI Setup
window = ctk.CTk()
window.title("🧠 Super Humanizer GUI (Fully Dynamic Models)")
window.geometry("700x700")

input_folder_var = ctk.StringVar()
output_folder_var = ctk.StringVar()

available_models = get_installed_ollama_models()
if available_models:
    model_var = ctk.StringVar(value=available_models[0])
else:
    model_var = ctk.StringVar(value="No models found")

grammar_var = ctk.BooleanVar()

# Input Folder
ctk.CTkLabel(window, text="Input Folder:", font=("Arial", 14)).pack(pady=(10, 0))
ctk.CTkEntry(window, textvariable=input_folder_var, width=500).pack()
ctk.CTkButton(window, text="Select Input Folder", command=lambda: input_folder_var.set(filedialog.askdirectory())).pack(pady=5)

# Output Folder
ctk.CTkLabel(window, text="Output Folder:", font=("Arial", 14)).pack(pady=(10, 0))
ctk.CTkEntry(window, textvariable=output_folder_var, width=500).pack()
ctk.CTkButton(window, text="Select Output Folder", command=lambda: output_folder_var.set(filedialog.askdirectory())).pack(pady=5)

# Model Selection
ctk.CTkLabel(window, text="Select Model:", font=("Arial", 14)).pack(pady=(10, 0))
ctk.CTkOptionMenu(window, variable=model_var, values=available_models if available_models else ["No models found"]).pack(pady=5)

# Grammar Correction Checkbox
ctk.CTkCheckBox(window, text="Enable Grammar Correction", variable=grammar_var).pack(pady=10)

# Start Button
ctk.CTkButton(window, text="🚀 Start Humanizing", command=start_batch, fg_color="green", hover_color="darkgreen", font=("Arial", 16)).pack(pady=10)

# Progress Bar
progress_bar = ctk.CTkProgressBar(window, width=500)
progress_bar.set(0)
progress_bar.pack(pady=10)

# Log Area
log_area = ctk.CTkTextbox(window, width=650, height=250)
log_area.pack(pady=10)

window.mainloop()
